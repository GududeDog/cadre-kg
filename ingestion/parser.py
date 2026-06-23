import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Union

from openpyxl import load_workbook


@dataclass
class Document:
    file_path: Path
    doc_type: str  # cadre_form / evaluation / investigation / report / work_division / txt
    raw_text: str = ""
    paras_text: str = ""  # paragraph-only text (no table | join)
    tables: List[Dict] = field(default_factory=list)
    metadata: Dict = field(default_factory=dict)
    sheets: Dict[str, List[Dict]] = field(default_factory=dict)


class DocumentParser:
    SUPPORTED = {".doc", ".docx", ".xls", ".xlsx", ".txt"}

    @classmethod
    def parse(cls, file_path: Union[str, Path]) -> Optional[Document]:
        path = Path(file_path)
        if not path.exists():
            return None
        suffix = path.suffix.lower()
        if suffix not in cls.SUPPORTED:
            return None
        handler = getattr(cls, f"_parse_{suffix[1:]}", None)
        if handler is None:
            return None
        doc = handler(path)
        doc.doc_type = cls._infer_type(path)
        # Special handling for cadre_form tables
        if doc.doc_type == "cadre_form":
            doc = cls.parse_cadre_form(doc)
        return doc

    # ── doc ──────────────────────────────────────────────
    @staticmethod
    def _parse_doc(path: Path) -> Document:
        """Convert .doc → .docx via win32com, then parse."""
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        abs_path = str(path.resolve())
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            docx_path = tmp.name
        try:
            wdoc = word.Documents.Open(abs_path)
            wdoc.SaveAs(docx_path, FileFormat=16)  # 16 = wdFormatXMLDocument
            wdoc.Close()
        finally:
            word.Quit()
        doc = DocumentParser._parse_docx(Path(docx_path))
        Path(docx_path).unlink(missing_ok=True)
        return doc

    # ── docx ─────────────────────────────────────────────
    @staticmethod
    def _parse_docx(path: Path) -> Document:
        from docx import Document as DocxDoc
        d = DocxDoc(str(path))
        doc = Document(file_path=path, doc_type="")
        paras = [p.text for p in d.paragraphs if p.text.strip()]
        doc.raw_text = "\n".join(paras)
        doc.paras_text = doc.raw_text  # save clean paragraph text
        for table in d.tables:
            raw_rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            rows = raw_rows[1:] if len(raw_rows) > 1 else []
            headers = raw_rows[0] if raw_rows else []
            doc.tables.append({"headers": headers, "rows": rows, "raw_rows": raw_rows})
            doc.raw_text += "\n" + "\n".join(" | ".join(r) for r in raw_rows)
        return doc

    # ── xlsx / xls ──────────────────────────────────────
    @staticmethod
    def _parse_xlsx(path: Path) -> Document:
        wb = load_workbook(str(path), data_only=True)
        doc = Document(file_path=path, doc_type="")
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            rows = [[str(c) if c is not None else "" for c in row] for row in rows]
            headers = rows[0] if rows else []
            data_rows = rows[1:] if len(rows) > 1 else []
            records = []
            for row in data_rows:
                records.append(dict(zip(headers, row)) if len(headers) == len(row) else row)
            doc.sheets[sheet_name] = records
            # raw_text
            doc.raw_text += f"\n【{sheet_name}】\n"
            doc.raw_text += "\n".join(" | ".join(row) for row in rows)
        return doc

    _parse_xls = _parse_xlsx  # openpyxl handles both

    # ── txt ──────────────────────────────────────────────
    @staticmethod
    def _parse_txt(path: Path) -> Document:
        raw = path.read_text(encoding="utf-8", errors="replace")
        return Document(file_path=path, doc_type="", raw_text=raw)

    # ── cadre_form 专用解析 ─────────────────────────────
    @staticmethod
    def parse_cadre_form(doc: Document) -> Document:
        """Clean cadre_form: use paragraph text + deduped table rows."""
        paras_text = doc.paras_text or ""

        table_lines = []
        for table in doc.tables:
            raw = table.get("raw_rows", [])
            if not raw:
                continue
            for row in raw:
                prev = None
                parts = []
                for cell in row:
                    cell = cell.replace("\n", " ").replace("\r", " ").strip()
                    if cell and cell != prev:
                        parts.append(cell)
                    prev = cell
                if parts:
                    table_lines.append("  ".join(parts))

        doc.raw_text = paras_text
        if table_lines:
            doc.raw_text += "\n" + "\n".join(table_lines)
        return doc

    # ── type inference ──────────────────────────────────
    _TYPE_MAP = {
        "干部审批表": "cadre_form",
        "任免表": "cadre_form",
        "民主测评汇总表": "evaluation_summary",
        "考察材料": "investigation",
        "述职报告": "report",
        "分工备案": "work_division",
        "民主测评结果": "evaluation_result",
        "领导班子民主测评": "evaluation_team",
        "领导班子分工备案表": "work_division_table",
        "功能需求": "requirements",
    }

    @classmethod
    def _infer_type(cls, path: Path) -> str:
        stem = path.stem
        for kw, t in cls._TYPE_MAP.items():
            if kw in stem:
                return t
        for parent in path.parents:
            parent_name = parent.name
            for kw, t in cls._TYPE_MAP.items():
                if kw in parent_name:
                    return t
        return "txt"
